"""组卷 — 删模板旧内容 → 填新题目"""
import json, os, random, shutil, re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from mkexam.omml import convert_text_to_omml, latex_to_omml, insert_omml_into_paragraph
from lxml import etree

BASE = os.path.dirname(os.path.abspath(__file__))
TPL_DIR = os.path.join(BASE, "试卷模板")
OUT_DIR = os.path.join(BASE, "output")

ANSWER_LABELS = ["A", "B", "C", "D", "E"]
PROPOSER = "（自动生成）"
REVIEWER = "（待审核）"
POLITICAL_REVIEW = "（待审核）"
YEAR = "2025"
DEPARTMENT = "信息工程学院"
MAJOR = "电子信息工程技术"
EXAM_TYPE = "理论"
EXAM_METHOD = "闭卷"
PAPER_CODE = "A卷"

nsmap = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
}

def remove_paragraph(para):
    """从文档中删除一个段落"""
    p = para._element
    p.getparent().remove(p)

def remove_all_content_after(doc, marker_text):
    """删除文档中从标记段落之后的所有内容（包括空段）"""
    started = False
    to_remove = []
    for para in doc.paragraphs:
        if marker_text and marker_text in para.text:
            started = True
            continue
        if started:
            to_remove.append(para)
    for p in to_remove:
        remove_paragraph(p)

def remove_paragraphs_containing(doc, keywords):
    """删除包含指定关键词的段落"""
    to_remove = []
    for para in doc.paragraphs:
        for kw in keywords:
            if kw in para.text:
                to_remove.append(para)
                break
    for p in to_remove:
        remove_paragraph(p)

def remove_table(doc, table_idx):
    """删除指定索引的表格"""
    if table_idx < len(doc.tables):
        tbl = doc.tables[table_idx]._tbl
        tbl.getparent().remove(tbl)

def main():
    os.makedirs(OUT_DIR, exist_ok=True)

    courses = [
        {
            "name": "电工电子技术",
            "json": os.path.join(BASE, "..", "electest", "电工电子技术", "questions.json"),
            "is_list": True,
            "sub_name": None,
            "sections": [
                ("一、单项选择题", "choice", 15, 2),
                ("二、判断题", "tf", 5, 2),
                ("三、填空题", "fill", 5, 2),
                ("四、计算分析题", "calc", 5, 10),
            ],
        },
        {
            "name": "单片机技术",
            "json": os.path.join(BASE, "..", "electest", "传感器单片机物联网", "subjects.json"),
            "is_list": False,
            "sub_name": "单片机技术",
            "sections": [
                ("一、单项选择题", "choice", 15, 2),
                ("二、判断题", "tf", 5, 2),
                ("三、填空题", "fill", 5, 4),
                ("四、简答题", "short", 2, 10),
                ("五、分析题", "analysis", 2, 10),
            ],
        },
    ]

    for course in courses:
        print(f"\n{'='*50}")
        print(f"{course['name']}")
        print(f"{'='*50}")

        # 加载题库
        if course["is_list"]:
            with open(course["json"]) as f:
                bank = json.load(f)
        else:
            with open(course["json"]) as f:
                subjects = json.load(f)
            bank = []
            for s in subjects:
                if course["sub_name"] and course["sub_name"] in s["name"]:
                    bank = s["questions"]
                    break
        print(f"题库: {len(bank)} 题")

        # 抽题
        selected = {}
        for title, key, count, score in course["sections"]:
            pool = [q for q in bank if q.get("type") == key and q.get("difficulty", 1) <= 2]
            if len(pool) < count:
                pool = [q for q in bank if q.get("type") == key]
            qs = random.sample(pool, min(count, len(pool)))
            selected[key] = qs
            print(f"  {title}: {len(qs)} 题")

        total = sum(len(selected.get(k, [])) * s for _, k, _, s in course["sections"])
        print(f"总分: {total} 分")
        ck = selected.get("choice", [])
        tf = selected.get("tf", [])
        print(f"  选择+判断总分: {len(ck)*2 + len(tf)*2} 分")

        stem = f"湖北职业技术学院_2025-2026学年第二学期_{course['name']}_期末考试A卷"

        # ═══════════ 试卷 ═══════════
        src = os.path.join(TPL_DIR, "2-1湖北职业技术学院2025-2026学年第二学期《课程名称》期末考试A卷（含答题纸）.docx")
        dst = os.path.join(OUT_DIR, f"{stem}.docx")
        shutil.copy(src, dst)
        doc = Document(dst)

        # 替换课名
        for para in doc.paragraphs:
            for run in para.runs:
                if "课程名称" in run.text:
                    run.text = run.text.replace("课程名称", course["name"])

        # 替换头部的占位符（整段替换，避免多 run 残留）
        for para in doc.paragraphs:
            full = para.text
            new_full = full
            if "课程名称" in full:
                new_full = full.replace("课程名称", course["name"])
            elif "考试课程：" in full and "XXXXX" in full:
                new_full = f"考试课程：{course['name']}          考试类型：{EXAM_TYPE}    考试方式：{EXAM_METHOD}    试卷：{PAPER_CODE}"
            elif "院系（专业）" in full:
                new_full = f"年级：{YEAR}    院系（专业）：{DEPARTMENT}{MAJOR}"
            elif "命题人：" in full and "XXXXX" in full:
                new_full = f"命题人：{PROPOSER}    审核人：{REVIEWER}    政审：{POLITICAL_REVIEW}"
            if new_full != full:
                # 清空所有 run 并设置新文本
                for run in para.runs:
                    run.text = ""
                para.runs[0].text = new_full

        # 删除所有题型表格（表1-7），只保留分数表（表0）
        # 注意：删除时从后往前删，避免索引变化
        for ti in range(len(doc.tables) - 1, 0, -1):
            tbl = doc.tables[ti]._tbl
            tbl.getparent().remove(tbl)

        # 删除注意事项之后的所有旧段落
        remove_all_content_after(doc, "注意事项")

        # 更新分数表（表0）
        if doc.tables:
            t0 = doc.tables[0]
            for ci, (title, key, count, score) in enumerate(course["sections"]):
                if ci >= 5: break
                total_sec = len(selected.get(key, [])) * score
                cn = title[0]  # 一、二、三、四
                t0.rows[0].cells[ci + 1].text = f"{cn}\n{total_sec}分"
            t0.rows[0].cells[6].text = str(total)

        # 图片目录（电工电子技术）
        ee_fig_dir = os.path.join(BASE, "..", "electest", "电工电子技术", "figures")

        # 追加新题目（每段之间空一行）
        for si, (title, key, count, score) in enumerate(course["sections"]):
            qs = selected.get(key, [])
            if not qs: continue
            total_sec = len(qs) * score
            if si > 0:
                doc.add_paragraph("")
            p = doc.add_paragraph()
            r = p.add_run(f"{title}（每题{score}分，共{total_sec}分）")
            r.bold = True
            r.font.size = Pt(12)
            iot_img_dir = os.path.join(BASE, "..", "electest", "传感器单片机物联网", "images")
            for i, q in enumerate(qs, 1):
                text = q.get("q", q.get("text", ""))
                img_path = None

                # 仅在计算/分析题中查图
                if key in ("calc", "analysis"):
                    m = re.search(r'(NEW-\d+)', text)
                    if m:
                        p = os.path.join(ee_fig_dir, f"图{m.group(1)}.png")
                        if os.path.exists(p):
                            img_path = p
                    if not img_path and key == "analysis":
                        akw = {'最小系统': 'm401_min_sys.png',
                               '晶振': 'm406_crystal.png',
                               'LED数码管': 'm402_led.png',
                               '矩阵键盘': 'm403_keypad.png',
                               '复位电路': 'm401_min_sys.png',
                               '光耦': 'm413_optocoupler.png',
                               '继电器': 'm414_relay_driver.png',
                               '电源电路': 'm415_power.png'}
                        for kw, fn in akw.items():
                            p = os.path.join(iot_img_dir, fn)
                            if kw in text and os.path.exists(p):
                                img_path = p
                                break

                # 先图后文（"如图 XX 所示"）
                if img_path:
                    try:
                        doc.add_picture(img_path, width=Inches(3.5))
                        doc.add_paragraph("")
                    except:
                        pass

                # 创建带 OMML 公式的段落
                p = doc.add_paragraph()
                r_num = p.add_run(f"{i}. ")
                segments = convert_text_to_omml(text)
                for seg_type, seg_content in segments:
                    if seg_type == 'text':
                        p.add_run(seg_content)
                    else:
                        processed, omml = latex_to_omml(seg_content)
                        if omml:
                            insert_omml_into_paragraph(p, omml)
                        else:
                            p.add_run(processed)

                if key == "choice":
                    opts = q.get("opts", q.get("options", []))
                    for oi, opt in enumerate(opts[:4]):
                        op = doc.add_paragraph()
                        op.add_run(f"    {ANSWER_LABELS[oi]}. ")
                        segs = convert_text_to_omml(opt)
                        for st, sc in segs:
                            if st == 'text':
                                op.add_run(sc)
                            else:
                                processed, omml = latex_to_omml(sc)
                                if omml:
                                    insert_omml_into_paragraph(op, omml)
                                else:
                                    op.add_run(processed)
                elif key == "tf":
                    doc.add_paragraph("    （  ）")
                else:
                    doc.add_paragraph("    ______")
                doc.add_paragraph("")

        doc.save(dst)
        print(f"  ✓ 试卷")

        # ═══════════ 评分标准 ═══════════
        src2 = os.path.join(TPL_DIR, "2-2湖北职业技术学院2025-2026学年第二学期《课程名称》期末考试A卷评分标准（含参考答案）.docx")
        dst2 = os.path.join(OUT_DIR, f"{stem}评分标准（含参考答案）.docx")
        shutil.copy(src2, dst2)
        doc2 = Document(dst2)

        # 替换头部的占位符（整段替换）
        for para in doc2.paragraphs:
            full = para.text
            new_full = full
            if "课程名称" in full:
                new_full = full.replace("课程名称", course["name"])
            elif "考试课程：" in full and "XXXXX" in full:
                new_full = f"考试课程：{course['name']}          考试类型：{EXAM_TYPE}    考试方式：{EXAM_METHOD}    试卷：{PAPER_CODE}"
            elif "院系（专业）" in full:
                new_full = f"年级：{YEAR}    院系（专业）：{DEPARTMENT}{MAJOR}"
            elif "命题人：" in full and "XXXXX" in full:
                new_full = f"命题人：{PROPOSER}    审核人：{REVIEWER}    政审：{POLITICAL_REVIEW}"
            if new_full != full:
                for run in para.runs:
                    run.text = ""
                para.runs[0].text = new_full

        # 删除旧的多选题和案例分析题相关内容
        remove_paragraphs_containing(doc2, ["多项选择题", "案例分析题", "旅游", "导游", "旅行社", "（4分）", "（2分）", "请分析本次事故", "发现只有38人", "已经误车的处理", "如何避免此类事故", "日程安排不当", "措施不当", "认真核实机"])

        # 删除多选表格（表1）
        if len(doc2.tables) > 1:
            remove_table(doc2, 1)

        # 更新单选题标题段落
        choice_qs = selected.get("choice", [])
        if choice_qs:
            for para in doc2.paragraphs:
                if "单项选择题" in para.text:
                    for run in para.runs:
                        run.text = f"一、单项选择题（共{len(choice_qs)}题，共{len(choice_qs)*2}分）"
                    break

        # 更新判断题标题段落
        tf_qs = selected.get("tf", [])
        if tf_qs:
            for para in doc2.paragraphs:
                if "判断题" in para.text:
                    for run in para.runs:
                        run.text = f"二、判断题（共{len(tf_qs)}题，共{len(tf_qs)*2}分）"
                    break

        # 填单选答案表（表0）
        if choice_qs and doc2.tables:
            t = doc2.tables[0]
            cols = len(t.columns)
            per_row = cols - 1
            for ri in range(min(4, len(t.rows))):
                for ci in range(1, min(len(t.rows[ri].cells), per_row + 1)):
                    idx = (ri // 2) * per_row + (ci - 1)
                    if idx >= len(choice_qs):
                        # 清空多余的单元格
                        t.rows[ri].cells[ci].text = ""
                        continue
                    if ri % 2 == 0:
                        t.rows[ri].cells[ci].text = str(idx + 1)
                    else:
                        ans = choice_qs[idx].get("ans", choice_qs[idx].get("answer", ""))
                        if isinstance(ans, int):
                            t.rows[ri].cells[ci].text = ANSWER_LABELS[ans] if ans < 5 else str(ans + 1)
                        else:
                            t.rows[ri].cells[ci].text = str(ans)

        # 填判断答案表（现在是表1，因为多选已被删除）
        if tf_qs and len(doc2.tables) > 1:
            t = doc2.tables[1]
            for ri in range(min(2, len(t.rows))):
                for ci in range(1, len(t.rows[ri].cells)):
                    idx = ci - 1
                    if idx >= len(tf_qs):
                        t.rows[ri].cells[ci].text = ""
                        continue
                    if ri % 2 == 0:
                        t.rows[ri].cells[ci].text = str(idx + 1)
                    else:
                        ans = tf_qs[idx].get("ans", tf_qs[idx].get("answer", ""))
                        if isinstance(ans, bool):
                            t.rows[ri].cells[ci].text = "√" if ans else "×"
                        else:
                            s = str(ans).strip().lower()
                            t.rows[ri].cells[ci].text = "√" if s in ("true", "正确") else "×"

        # 追加填空、计算/简答的答案
        doc2.add_paragraph("")
        for title, key, count, score in course["sections"]:
            qs = selected.get(key, [])
            if key in ("choice", "tf") or not qs:
                continue
            p = doc2.add_paragraph()
            r = p.add_run(f"{title}答案")
            r.bold = True
            for i, q in enumerate(qs, 1):
                ans = q.get("ans", q.get("answer", ""))
                doc2.add_paragraph(f"  {i}. {ans}")

        doc2.save(dst2)
        print(f"  ✓ 评分标准")

    print(f"\n输出: {OUT_DIR}/")
    print("完成")


if __name__ == "__main__":
    main()

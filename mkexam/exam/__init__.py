"""组卷引擎 — 选题 + 预览调题 + 生成试卷 + 批量出卷"""
import os, random, shutil, re, json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from mkexam.bank import BankManager
from mkexam.exam.preset import load_preset, save_preset, list_presets

_BASE = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
TPL_DIR = os.path.join(_BASE, "试卷模板")
FIG_DIR_EE = os.path.join(_BASE, "..", "electest", "电工电子技术", "figures")
FIG_DIR_IOT = os.path.join(_BASE, "..", "electest", "传感器单片机物联网", "images")


class ExamSelector:
    """选题引擎"""

    def __init__(self, bank: BankManager):
        self.bank = bank

    def auto_select(self, sub_name: str, sections: list, max_diff: int = 2) -> dict:
        """自动选题 — 相同图片的题目在跨题型间不重复"""
        questions = self.bank.list_questions(sub_name)
        selected = {}
        used_images: set[str] = set()  # 已选图片，跨题型不重复

        for title, key, count, score in sections:
            pool = [q for q in questions if q.get("type") == key and q.get("difficulty", 1) <= max_diff]
            if len(pool) < count:
                pool = [q for q in questions if q.get("type") == key]

            # 过滤掉图片已被其他题型使用的题目
            img_filtered = [q for q in pool
                            if not q.get("image") or q["image"] not in used_images]
            if len(img_filtered) >= count:
                pool = img_filtered

            # 同题型内也按图片去重（相同图片只选一题）
            from collections import defaultdict
            by_img: dict[str, list] = defaultdict(list)
            no_img_qs: list = []
            for q in pool:
                img = q.get("image")
                if img:
                    by_img[img].append(q)
                else:
                    no_img_qs.append(q)

            random.shuffle(no_img_qs)
            selected_qs: list = []
            # 每张图最多选一题
            for img, candidates in by_img.items():
                if len(selected_qs) >= count:
                    break
                selected_qs.append(random.choice(candidates))
            # 不够时用无图片的题目补
            while len(selected_qs) < count and no_img_qs:
                selected_qs.append(no_img_qs.pop())

            qs = selected_qs[:count]

            # 选择题/判断题：答案离散分布，不扎堆
            if key == "choice" and count >= 3:
                # 选项已打乱，直接随机抽即可保证分布
                random.shuffle(pool)
                qs = pool[:count]

            elif key == "tf" and count >= 2:
                def _is_true(q):
                    a = q.get("ans", q.get("answer", ""))
                    return a is True or str(a).strip().lower() in ("true", "正确", "√")
                def _is_false(q):
                    a = q.get("ans", q.get("answer", ""))
                    return a is False or str(a).strip().lower() in ("false", "错误", "×")
                has_true = any(_is_true(q) for q in qs)
                has_false = any(_is_false(q) for q in qs)
                retry = 0
                while (not has_true or not has_false) and retry < 10 and len(pool) >= count:
                    random.shuffle(pool)
                    qs = pool[:count]
                    has_true = any(q.get("answer") is True or str(q.get("answer", "")).lower() == "true" for q in qs)
                    has_false = any(q.get("answer") is False or str(q.get("answer", "")).lower() == "false" for q in qs)
                    retry += 1

            # 记录本次选中题的图片
            for q in qs:
                img = q.get("image")
                if img:
                    used_images.add(img)
            selected[key] = qs
        return selected

    def preview(self, selected: dict, detail: bool = False):
        """预览"""
        for key, qs in selected.items():
            print(f"  {key}: {len(qs)} 题")
            for i, q in enumerate(qs, 1):
                text = q.get("q", q.get("text", ""))
                print(f"    {i}. {text[:60]}")
                if detail and q.get("opts"):
                    for oi, opt in enumerate(q.get("opts", [])[:4]):
                        print(f"       {['A','B','C','D'][oi]}. {opt[:40]}")
            print()

    # ─── 手动调整 ────────────────────────────────

    def interactive_adjust(self, sub_name: str, selected: dict, sections: list) -> dict:
        """交互式预览+调整选题"""
        while True:
            print("\n" + "=" * 50)
            print("  当前选题预览")
            print("=" * 50)
            self.preview(selected, detail=True)

            total = sum(len(selected.get(k, [])) * s for _, k, _, s in sections)
            print(f"  总分: {total} 分")

            print("\n  操作:")
            print("    r <题型>  — 重新抽取该题型（如: r choice）")
            print("    s        — 保存并继续")
            print("    q        — 取消")

            cmd = input("  > ").strip()
            if cmd.lower() == "q":
                return None
            if cmd.lower() == "s":
                return selected
            if cmd.startswith("r "):
                key = cmd[2:].strip()
                # 找到对应的配置
                for title, k, count, score in sections:
                    if k == key:
                        questions = self.bank.list_questions(sub_name)
                        pool = [q for q in questions if q.get("type") == k]
                        qs = random.sample(pool, min(count, len(pool)))
                        selected[k] = qs
                        print(f"  ✅ {title} 已重新抽取 {len(qs)} 题")
                        break
                else:
                    print(f"  未知题型: {key}")

    # ─── 批量选题 ────────────────────────────────

    def batch_select(self, sub_name: str, sections: list, n_sets: int) -> list[dict]:
        """生成多套选题（各套题目不重复 + 同卷同图不跨题型）"""
        all_questions = self.bank.list_questions(sub_name)
        results = []
        used_ids = set()

        for set_i in range(n_sets):
            selected = {}
            used_images: set[str] = set()
            for title, key, count, score in sections:
                pool = [q for q in all_questions if q.get("type") == key
                        and q.get("id") not in used_ids]
                if len(pool) < count:
                    pool = [q for q in all_questions if q.get("type") == key]
                # 过滤掉图片已被同一卷其他题型使用的题目
                img_filtered = [q for q in pool
                                if not q.get("image") or q["image"] not in used_images]
                if len(img_filtered) >= count:
                    pool = img_filtered
                # 同题型内也按图片去重
                from collections import defaultdict
                by_img: dict[str, list] = defaultdict(list)
                no_img_qs: list = []
                for q in pool:
                    img = q.get("image")
                    if img:
                        by_img[img].append(q)
                    else:
                        no_img_qs.append(q)
                random.shuffle(no_img_qs)
                selected_qs: list = []
                for img, candidates in by_img.items():
                    if len(selected_qs) >= count:
                        break
                    selected_qs.append(random.choice(candidates))
                while len(selected_qs) < count and no_img_qs:
                    selected_qs.append(no_img_qs.pop())
                qs = selected_qs[:count]

                # 选择题/判断题答案离散
                if key == "choice" and count >= 3:
                    # 选项已打乱，直接随机抽
                    random.shuffle(pool)
                    qs = pool[:count]
                elif key == "tf" and count >= 2:
                    def _is_true(q):
                        a = q.get("ans", q.get("answer", ""))
                        return a is True or str(a).strip().lower() in ("true", "正确", "√")
                    def _is_false(q):
                        a = q.get("ans", q.get("answer", ""))
                        return a is False or str(a).strip().lower() in ("false", "错误", "×")
                    pool_t = [q for q in pool if _is_true(q)]
                    pool_f = [q for q in pool if _is_false(q)]
                    random.shuffle(pool_t); random.shuffle(pool_f)
                    # 尽量均匀分布，不够则全部从另一pool补
                    n_true = min(len(pool_t), max(1, count // 2))
                    n_false = min(len(pool_f), count - n_true)
                    if n_true + n_false < count:
                        # 补足剩余
                        if len(pool_t) > n_true:
                            n_true = min(len(pool_t), count - n_false)
                        elif len(pool_f) > n_false:
                            n_false = min(len(pool_f), count - n_true)
                    qs = pool_t[:n_true] + pool_f[:n_false]
                    random.shuffle(qs)

                for q in qs:
                    used_ids.add(q.get("id", ""))
                    img = q.get("image")
                    if img:
                        used_images.add(img)
                selected[key] = qs
            results.append(selected)
        return results


class PaperGenerator:
    """试卷生成器"""

    def __init__(self, output_dir: str):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        self.header_info: dict = {}
        self.has_answer_sheet: bool = True
        self.current_label: str = ""

    def generate(self, sub_name: str, selected: dict, sections: list, label: str = ""):
        """生成一份试卷和评分标准"""
        course_map = {"单片机技术": "单片机技术应用", "电工电子技术": "电工电子技术"}
        display_name = course_map.get(sub_name, sub_name)
        suffix = f"{label}卷" if label else ""
        stem = f"湖北职业技术学院2025-2026学年第二学期《{display_name}》期末考试{suffix}"
        self.current_label = label
        # 按科目建子目录
        sub_dir = os.path.join(self.output_dir, display_name)
        os.makedirs(sub_dir, exist_ok=True)
        old_dir = self.output_dir
        self.output_dir = sub_dir
        self._make_paper(sub_name, selected, sections, stem)
        if self.has_answer_sheet:
            self._make_answer_sheet(sub_name, selected, sections, stem)
        self._make_scoring(sub_name, selected, sections, stem)
        self.output_dir = old_dir  # 恢复
        print(f"  ✓ {display_name}/{stem}")

    def batch_generate(self, sub_name: str, selected_list: list[dict], sections: list, labels: list[str] = None):
        """批量生成多份试卷"""
        if labels is None:
            labels = [chr(65 + i) for i in range(len(selected_list))]  # A, B, C...
        for i, selected in enumerate(selected_list):
            self.generate(sub_name, selected, sections, labels[i] if i < len(labels) else str(i + 1))
        print(f"\n✅ 共生成 {len(selected_list)} 份试卷")

    # ─── 试卷生成 ────────────────────────────────

    def _make_paper(self, sub_name, selected, sections, stem):
        if self.has_answer_sheet:
            tpl = "3-1湖北职业技术学院2025-2026学年第二学期《课程名称》期末考试A卷.docx"
        else:
            tpl = "2-1湖北职业技术学院2025-2026学年第二学期《课程名称》期末考试A卷（含答题纸）.docx"
        src = os.path.join(TPL_DIR, tpl)
        if os.path.exists(src):
            dst = os.path.join(self.output_dir, f"{stem}.docx")
            shutil.copy(src, dst)
            doc = Document(dst)
        else:
            doc = Document()

        self._replace_placeholders(doc, sub_name)
        self._clean_template(doc, sections, selected)
        self._append_questions(doc, selected, sections)

        doc.save(os.path.join(self.output_dir, f"{stem}.docx"))

    def _make_answer_sheet(self, sub_name, selected, sections, stem):
        """生成独立答题纸"""
        src = os.path.join(TPL_DIR, "3-2湖北职业技术学院2025-2026学年第二学期《课程名称》期末考试A卷答题纸.docx")
        dst = os.path.join(self.output_dir, f"{stem}答题纸.docx")
        if os.path.exists(src):
            shutil.copy(src, dst)
            doc = Document(dst)
        else:
            doc = Document()
        self._replace_placeholders(doc, sub_name)

        # 清除头部后的旧内容（含大量空白段）
        to_remove = []
        started = False
        for para in doc.paragraphs:
            if '说明' in para.text or '注意事项' in para.text:
                started = True; continue
            if started:
                to_remove.append(para)
        for p in to_remove:
            try: p._element.getparent().remove(p._element)
            except: pass
        # 保留第一个表格（得分汇总表）用于动态更新，删除其余旧表格
        score_table = doc.tables[0] if doc.tables else None
        for ti in range(len(doc.tables) - 1, 0, -1):
            try: doc.tables[ti]._tbl.getparent().remove(doc.tables[ti]._tbl)
            except: pass

        SECTION_DESCS = {
            "choice": "从每小题给出的四个备选项中选出符合题目要求的一项。", "multiple": "从每小题给出的备选项中选出符合题目要求的选项，多选、少选、错选均不得分。",
            "tf": "判断以下说法是否正确，正确的打\"√\"，错误的打\"×\"。",
            "fill": "将正确答案填写在横线上。",
            "short": "简要回答下列问题。",
            "calc": "写出必要的分析计算过程。",
            "分析题": "分析下列问题并作答。",
            "应用题": "根据给定条件完成程序设计。",
            "应用分析题": "分析下列问题并作答。",
        }
        KEY_TITLES = {
            "choice": "选择题", "multiple": "多选题", "tf": "判断题", "fill": "填空题",
            "short": "简答题", "calc": "计算分析题",
            "分析题": "分析题", "应用题": "应用题", "应用分析题": "应用分析题",
        }
        CN_NUMS = ["一", "二", "三", "四", "五", "六", "七", "八"]
        q_global = 0

        def _keep_with_next(para):
            """防止段落与其后续内容跨页断开"""
            pPr = para._p.get_or_add_pPr()
            kn = OxmlElement('w:keepNext')
            pPr.append(kn)

        for si, (title, key, count, score) in enumerate(sections):
            qs = selected.get(key, [])
            if not qs: continue
            total_sec = len(qs) * score
            if not title.strip():
                cn = CN_NUMS[si] if si < len(CN_NUMS) else str(si + 1)
                type_name = KEY_TITLES.get(key, key)
                title = f"{cn}、{type_name}"

            p = doc.add_paragraph()
            p.add_run(f"{title}（每题{score}分，共{total_sec}分）").bold = True
            desc = SECTION_DESCS.get(key, "")
            if desc:
                dr = p.add_run(f"  {desc}")
                dr.bold = True

            if key in ("choice", "tf"):
                cols = 5
                n_rows = (len(qs) + cols - 1) // cols
                table = doc.add_table(rows=n_rows * 2, cols=cols)
                from docx.oxml.ns import qn
                # 设置紧凑列宽（缩约40%）
                for row in table.rows:
                    for cell in row.cells:
                        cell.width = Inches(0.7)
                tbl_pr = table._tbl.tblPr if table._tbl.tblPr is not None else OxmlElement('w:tblPr')
                tbl_borders = OxmlElement('w:tblBorders')
                for edge in ('top','left','bottom','right','insideH','insideV'):
                    border = OxmlElement(f'w:{edge}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')
                    tbl_borders.append(border)
                tbl_pr.append(tbl_borders)
                for qi in range(len(qs)):
                    q_global += 1
                    ri = (qi // cols) * 2; ci = qi % cols
                    table.rows[ri].cells[ci].text = str(q_global)
                    # 题号居中
                    for pa in table.rows[ri].cells[ci].paragraphs:
                        pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    table.rows[ri + 1].cells[ci].text = ""
            elif key == "fill":
                # 填空题：每行3空
                per_row = 3
                for row_start in range(0, len(qs), per_row):
                    bp = doc.add_paragraph()
                    _keep_with_next(bp)
                    row_qs = qs[row_start:row_start + per_row]
                    for qi, q in enumerate(row_qs):
                        q_global += 1
                        if qi > 0:
                            bp.add_run("  ")
                        bp.add_run(f"{q_global}. ")
                        r = bp.add_run(' ' * 20)
                        r.underline = True
            else:
                blank_lines = 10 if key in ("short", "calc") else (10 if key in ("分析题", "应用题") else 4)
                sub_blank = 10 if key in ("应用分析题", "分析题", "应用题") else 0
                for qi, q in enumerate(qs):
                    q_global += 1
                    text = q.get("q", q.get("text", ""))
                    sub_qs = re.findall(r'(?:^|\n)\s*（(\d+)）', text)
                    
                    bp = doc.add_paragraph()
                    _keep_with_next(bp)
                    run_text = f"{q_global}. "
                    
                    if sub_qs and key in ("应用分析题", "分析题", "应用题"):
                        run_text += f"（1）"
                        bp.add_run(run_text)
                        sub_lines = sub_blank if sub_blank else (blank_lines // len(sub_qs))
                        for _ in range(sub_lines):
                            ep = doc.add_paragraph("")
                            _keep_with_next(ep)
                        for si in range(1, len(sub_qs)):
                            doc.add_paragraph("")
                            sp = doc.add_paragraph()
                            _keep_with_next(sp)
                            sp.add_run(f"  （{si+1}）")
                            for _ in range(sub_lines):
                                ep = doc.add_paragraph("")
                                _keep_with_next(ep)
                    else:
                        bp.add_run(run_text)
                        for _ in range(blank_lines):
                            ep = doc.add_paragraph("")
                            _keep_with_next(ep)

        # 更新得分汇总表（只填题型和总分，得分留空供阅卷手写）
        if score_table is not None and sections:
            try:
                row0 = score_table.rows[0]
                row1 = score_table.rows[1]
                cn = ['一','二','三','四','五','六','七','八']
                total_score = 0
                # 填题型名称和合计分数，清空得分格
                for ci, (title, key, count, score) in enumerate(sections):
                    col = ci + 1
                    if col >= len(row0.cells) - 1: break
                    sec_score = len(selected.get(key, [])) * score
                    total_score += sec_score
                    row0.cells[col].text = cn[ci] if ci < len(cn) else str(ci+1)
                    row1.cells[col].text = ''  # 得分空着
                # 清空剩余列
                for ci in range(len(sections), len(row0.cells) - 2):
                    row0.cells[ci + 1].text = ''
                # 总分列（总分标在第一行，第二行留空）
                if len(row0.cells) >= len(sections) + 2:
                    row0.cells[len(sections) + 1].text = '总分'
                    row1.cells[len(sections) + 1].text = ''
                # 所有单元格居中
                for row in score_table.rows:
                    for cell in row.cells:
                        for pa in cell.paragraphs:
                            pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                print(f"  更新得分表失败: {e}")

        doc.save(dst)

    def _make_scoring(self, sub_name, selected, sections, stem):
        if self.has_answer_sheet:
            tpl = "3-3湖北职业技术学院2025-2026学年第二学期《课程名称》期末考试A卷评分标准（含参考答案）.docx"
        else:
            tpl = "2-2湖北职业技术学院2025-2026学年第二学期《课程名称》期末考试A卷评分标准（含参考答案）.docx"
        src = os.path.join(TPL_DIR, tpl)
        if os.path.exists(src):
            dst = os.path.join(self.output_dir, f"{stem}评分标准.docx")
            shutil.copy(src, dst)
            doc = Document(dst)
        else:
            doc = Document()

        self._replace_placeholders(doc, sub_name)
        # 清除旧模板内容，从零构建答案部分
        self._clean_scoring_template(doc, selected, sections)
        # 在「湖北职业技术学院」标题后插入副标题
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.shared import Pt
        for para in doc.paragraphs:
            if '湖北职业技术学院' in para.text:
                new_p = OxmlElement('w:p')
                pPr = OxmlElement('w:pPr')
                jc = OxmlElement('w:jc')
                jc.set(qn('w:val'), 'center')
                pPr.append(jc)
                new_p.append(pPr)
                r_elem = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:eastAsia'), '黑体')
                rFonts.set(qn('w:ascii'), '黑体')
                rPr.append(rFonts)
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), '32')
                rPr.append(sz)
                szCs = OxmlElement('w:szCs')
                szCs.set(qn('w:val'), '32')
                rPr.append(szCs)
                r_elem.append(rPr)
                t_elem = OxmlElement('w:t')
                t_elem.text = '评分标准（含参考答案）'
                r_elem.append(t_elem)
                new_p.append(r_elem)
                para._p.addnext(new_p)
                break
        self._build_scoring_sections(doc, selected, sections)

        doc.save(os.path.join(self.output_dir, f"{stem}评分标准.docx"))

    # ─── 辅助方法 ────────────────────────────────

    def _replace_placeholders(self, doc, sub_name):
        h = self.header_info
        # 替换课程名称（python-docx方式）
        for para in doc.paragraphs:
            for run in para.runs:
                if "课程名称" in run.text:
                    run.text = run.text.replace("课程名称", sub_name)
        # 替换 A卷/B卷 标记
        if self.current_label:
            vol = f"{self.current_label}卷"
            from lxml import etree
            body = doc.element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}body')
            if body is None:
                body = doc.element
            for elem in body.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                if elem.text:
                    t = elem.text.strip()
                    if t in ('A卷', 'A 卷'):
                        elem.text = vol
                    elif t in ('B卷', 'B 卷'):
                        elem.text = vol
        # 用lxml替换占位符（解决label和value在不同run的问题）
        if h.get("course") or h.get("author") or h.get("year") or h.get("dept"):
            try:
                from lxml import etree
                body = doc.element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}body')
                if body is None:
                    body = doc.element
                last_label = ''
                for elem in body.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                    if elem.text is None: continue
                    t = elem.text.strip()
                    if '考试课程' in t: last_label = 'course'
                    elif '命题人' in t: last_label = 'author'
                    elif '审核人' in t: last_label = 'reviewer'
                    elif '政审' in t: last_label = 'political'
                    elif '院系' in t: last_label = 'dept'
                    if t == 'XXXXX':
                        pad = 6
                        if last_label == 'course':
                            pad = 4
                        elif last_label in ('author', 'reviewer', 'political'):
                            pad = 8
                        val = ''
                        if last_label == 'course' and h.get('course'):
                            val = h['course']
                        elif last_label == 'author' and h.get('author'):
                            val = h['author']
                        elif last_label == 'reviewer' and h.get('reviewer'):
                            val = h['reviewer']
                        elif last_label == 'political' and h.get('political'):
                            val = h['political']
                        if val:
                            elem.text = ' ' * pad + val + ' ' * pad
                    elif 'XXXXX学院XXXXX专业' in elem.text and h.get('dept'):
                        raw = elem.text
                        elem.text = raw.replace('XXXXX学院XXXXX专业', h['dept'])
                    elif '202X' in elem.text and h.get('year'):
                        raw = elem.text
                        elem.text = raw.replace('202X', h['year'])
                    # 统一清理标签周围的空格
                    if '院系' in elem.text or '命题人' in elem.text or '审核人' in elem.text or '政审' in elem.text or '考试课程' in elem.text:
                        elem.text = elem.text.strip()
                    if elem.text and elem.text.strip() in ('A卷', 'A 卷', 'B卷', 'B 卷'):
                        elem.text = elem.text.strip()
            except Exception as e:
                print(f"  头部替换失败: {e}")

    def _clean_template(self, doc, sections, selected):
        # 删除注意事项后的段落
        started = False
        to_remove = []
        for para in doc.paragraphs:
            if "注意事项" in para.text:
                started = True
                continue
            if started:
                to_remove.append(para)
        for p in to_remove:
            try:
                p._element.getparent().remove(p._element)
            except:
                pass
        # 删除多余表格（保留表0）
        for ti in range(len(doc.tables) - 1, 0, -1):
            try:
                doc.tables[ti]._tbl.getparent().remove(doc.tables[ti]._tbl)
            except:
                pass
        # 更新分数表
        if doc.tables:
            t0 = doc.tables[0]
            cn = ["一", "二", "三", "四", "五", "六"]
            for ci, (title, key, count, score) in enumerate(sections):
                if ci >= 6: break
                total_sec = len(selected.get(key, [])) * score
                t0.rows[0].cells[ci + 1].text = f"{cn[ci]}\n{total_sec} 分"
            total = sum(len(selected.get(k, [])) * s for _, k, _, s in sections)
            t0.rows[0].cells[6].text = str(total)

    def _append_questions(self, doc, selected, sections):
        # 大题标题映射（title为空时自动生成）
        KEY_TITLES = {
            "choice": "选择题", "multiple": "多选题", "tf": "判断题", "fill": "填空题",
            "short": "简答题", "calc": "计算分析题", "analysis": "分析题",
            "分析题": "分析题", "应用题": "应用题",
        }
        CN_NUMS = ["一", "二", "三", "四", "五", "六", "七", "八"]
        # 大题描述文字
        SECTION_DESCS = {
            "choice": "从每小题给出的四个备选项中选出符合题目要求的一项，并将答案填在答题纸对应题号中。多选、少选、错选均不得分。", "multiple": "从每小题给出的备选项中选出符合题目要求的选项，多选、少选、错选均不得分。",
            "tf": "判断以下说法是否正确，正确的在括号内打\"√\"，错误的打\"×\"。",
            "fill": "将正确答案填写在横线上，每空答案唯一。",
            "short": "简要回答下列问题，要点明确，表述完整。",
            "calc": "写出必要的分析计算过程，只写结果不得分。",
            "analysis": "分析下列问题并作答，要求逻辑清晰、表述完整。",
            "分析题": "分析下列问题并作答，要求逻辑清晰、表述完整。",
            "应用题": "根据给定条件完成程序设计或代码填空，代码应符合C51语法规范。",
            "应用分析题": "应用分析下列问题并作答，要求逻辑清晰、表述完整。",
        }

        q_abc = 0  # 选择/填空/判断连续编号
        for si, (title, key, count, score) in enumerate(sections):
            qs = selected.get(key, [])
            if not qs: continue
            total_sec = len(qs) * score
            # title 为空时自动生成 "一、选择题" 格式
            if not title.strip():
                cn = CN_NUMS[si] if si < len(CN_NUMS) else str(si + 1)
                type_name = KEY_TITLES.get(key, key)
                title = f"{cn}、{type_name}"
            p = doc.add_paragraph()
            r = p.add_run(f"{title}（每题{score}分，共{total_sec}分）")
            r.bold = True
            r.font.name = '黑体'
            r._r.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            r.font.size = Pt(12)  # 小四号

            # 大题描述（跟标题同一行，统一粗体）
            desc = SECTION_DESCS.get(key, "")
            if desc:
                dr = p.add_run(f"\n{desc}")
                dr.bold = True
                dr.font.name = '黑体'
                dr._r.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

            # 试题卷不显示选择题答案表格（答题纸中已包含）

            for q in qs:
                if key in ("choice", "tf", "fill"):
                    q_abc += 1
                    display_q = q_abc
                else:
                    # 大题内每道题递增编号（大题重新从1开始）
                    display_q = qs.index(q) + 1
                text = q.get("q", q.get("text", ""))
                # 去掉 【图X·题型】 前缀
                text = re.sub(r'^【[^】]+】\s*', '', text)
                score_label = f"（{score}分）" if key not in ("choice", "tf", "fill") else ""
                # 多小问：拆分为独立段落（每小问独立一行）
                sub_segments = [(text, score_label)]
                if key not in ("choice", "tf", "fill"):
                    sub_qs = re.findall(r'(?:^|\n)\s*（(\d+)）', text)
                    if sub_qs:
                        n = len(sub_qs)
                        exp = q.get("exp", "")
                        scores_per_q = []
                        if exp:
                            parts = re.findall(r'（\d+）\s*(\d+)\s*分', exp)
                            if parts and len(parts) == n:
                                scores_per_q = [int(p) for p in parts]
                        if not scores_per_q:
                            base = score // n
                            remain = score % n
                            scores_per_q = [base + (1 if i < remain else 0) for i in range(n)]
                        # 提取题干前缀（第一个小问之前的内容）
                        first_m = re.search(r'(?:^|\n)\s*（(\d+)）', text)
                        if first_m:
                            if first_m.start() > 0:
                                prefix = text[:first_m.start()].strip()
                                sub_segments = [(prefix, '')] if prefix else []
                            else:
                                sub_segments = []
                        else:
                            sub_segments = []
                        # 逐小问拆分：从第一个小问开始
                        remaining = text if first_m is None else text
                        for si in range(n):
                            m = re.search(r'(?:^|\n)(\s*（(\d+)）)', remaining)
                            if not m: break
                            after_marker = remaining[m.end():]
                            next_m = re.search(r'\n\s*（(\d+)）', after_marker)
                            if next_m:
                                q_body = after_marker[:next_m.start()].strip()
                            else:
                                q_body = after_marker.strip()
                            sub_segments.append((f"（{m.group(2)}）{q_body}", scores_per_q[si]))
                            if next_m:
                                remaining = after_marker[next_m.start():]
                            else:
                                break
                p_q = doc.add_paragraph()
                p_q.add_run(f"{display_q}. ")
                # 渲染每个段落（子问题与题干紧密排列）
                for seg_i, (seg_text, seg_score) in enumerate(sub_segments):
                    if seg_i > 0:
                        sp = doc.add_paragraph()
                    else:
                        sp = p_q
                    self._add_formatted_text(sp, seg_text)
                    if isinstance(seg_score, int):
                        sp.add_run(f"  （{seg_score}分）")
                    elif seg_score:
                        sp.add_run(f"  {seg_score}")
                if key in ("choice", "multiple"):
                    opts = q.get("opts", q.get("options", []))
                    # 统一制表符对齐（所有选择题选项用同一位置）
                    def _add_tab(para):
                        pPr = para._p.get_or_add_pPr()
                        tabs = OxmlElement('w:tabs')
                        tab = OxmlElement('w:tab')
                        tab.set(qn('w:val'), 'left')
                        tab.set(qn('w:pos'), '4422')  # 20%页宽
                        tabs.append(tab)
                        pPr.append(tabs)
                    def _add_tab_run(para):
                        """在段落中添加一个制表符元素"""
                        r = OxmlElement('w:r')
                        tab_el = OxmlElement('w:tab')
                        r.append(tab_el)
                        para._p.append(r)
                    # A+B 一行
                    op = doc.add_paragraph()
                    _add_tab(op)
                    op.add_run("A. ")
                    self._add_formatted_text(op, opts[0])
                    _add_tab_run(op)
                    op.add_run("B. ")
                    self._add_formatted_text(op, opts[1])
                    # C+D 一行
                    cp = doc.add_paragraph()
                    _add_tab(cp)
                    cp.add_run("C. ")
                    self._add_formatted_text(cp, opts[2])
                    _add_tab_run(cp)
                    cp.add_run("D. ")
                    self._add_formatted_text(cp, opts[3])
                if key == "choice":
                    pass  # 选项已在上方处理
                elif key == "tf":
                    # 括号跟在题目后面，不换行
                    p_q.add_run("  （  ）")
                elif key == "fill":
                    pass  # 填空题题干已含下划线，不再追加
                # 图在题目下方（仅作图题）
                if key in ("分析题", "应用题", "应用分析题", "calc"):
                    self._insert_figure(doc, q)

    def _add_formatted_text(self, para, text):
        """将含 $...$ 的文本分段写入段落，公式部分转 Unicode"""
        from mkexam.omml import convert_text_to_omml, latex_to_omml, insert_omml_into_paragraph
        segments = convert_text_to_omml(text)
        for st, sc in segments:
            if st == 'text':
                para.add_run(sc)
            else:
                processed, omml = latex_to_omml(sc)
                if omml:
                    insert_omml_into_paragraph(para, omml)
                else:
                    para.add_run(processed)

    def _fill_answer_tables(self, doc, selected):
        tables = doc.tables
        choice_qs = selected.get("choice", [])
        if choice_qs and tables:
            t = tables[0]
            for ri in range(min(4, len(t.rows))):
                for ci in range(1, min(len(t.rows[ri].cells), 11)):
                    idx = (ri // 2) * 10 + (ci - 1)
                    if idx >= len(choice_qs):
                        t.rows[ri].cells[ci].text = ""
                        continue
                    if ri % 2 == 1:
                        ans = choice_qs[idx].get("ans", choice_qs[idx].get("answer", ""))
                        if isinstance(ans, int):
                            t.rows[ri].cells[ci].text = ["A","B","C","D"][ans] if ans < 4 else str(ans+1)
                        else:
                            t.rows[ri].cells[ci].text = str(ans)

        tf_qs = selected.get("tf", [])
        if tf_qs and len(tables) > 1:
            t = tables[1]
            for ri in range(min(2, len(t.rows))):
                for ci in range(1, len(t.rows[ri].cells)):
                    idx = ci - 1
                    if idx >= len(tf_qs):
                        t.rows[ri].cells[ci].text = ""
                        continue
                    if ri % 2 == 1:
                        ans = tf_qs[idx].get("ans", tf_qs[idx].get("answer", ""))
                        if isinstance(ans, bool):
                            t.rows[ri].cells[ci].text = "√" if ans else "×"
                        else:
                            t.rows[ri].cells[ci].text = "√" if str(ans).strip().lower() in ("true", "正确") else "×"

        if len(tables) > 1:
            for ri in range(len(tables[1].rows)):
                for ci in range(1, len(tables[1].columns)):
                    tables[1].rows[ri].cells[ci].text = ""

    def _append_other_answers(self, doc, selected, sections):
        for title, key, count, score in sections:
            qs = selected.get(key, [])
            if key in ("choice", "tf") or not qs: continue
            doc.add_paragraph("")
            p = doc.add_paragraph()
            r = p.add_run(f"{title}答案")
            r.bold = True
            for i, q in enumerate(qs, 1):
                ans = q.get("ans", q.get("answer", ""))
                steps = self._format_answer_steps(ans, score)
                doc.add_paragraph(f"  {i}.（共{score}分）")
                for step_text in steps:
                    doc.add_paragraph(f"    {step_text}")

    def _format_answer_steps(self, ans: str, total_score: int) -> list[str]:
        """将答案文本拆分为得分点列表"""
        if not ans:
            return ["见解析"]
        # 按分号或句号分割
        import re
        parts = re.split(r'[；;]\s*', ans)
        if len(parts) <= 1:
            parts = re.split(r'[。]\s*', ans)
        parts = [p.strip() for p in parts if p.strip()]
        if len(parts) <= 1:
            # 尝试按逗号分割（仅对计算题）
            if total_score >= 6:
                parts = re.split(r'[，,]\s*', ans)
                parts = [p.strip() for p in parts if p.strip()]

        if len(parts) <= 1:
            # 就一段话，直接返回
            return [f"{ans}（{total_score}分）"]

        # 均分分值
        base = total_score // len(parts)
        remain = total_score % len(parts)
        result = []
        for j, part in enumerate(parts):
            pts = base + (1 if j < remain else 0)
            result.append(f"{part}（{pts}分）")
        return result

    def _clean_scoring_template(self, doc, selected, sections):
        """清除评分标准模板中的旧内容"""
        # 旧模板的固定章节标题关键字（全部要删）
        old_headers = ["一、单项选择", "二、多项选择", "三、判断", "四、案例分析",
                       "五、", "六、", "七、", "八、",
                       "①", "②", "③", "④", "⑤",
                       "1.请分析", "2.发现", "3.已经", "4.如何避免",
                       "答案", "题号"]
        old_keywords = ["旅游", "导游", "旅行社",
                        "本次事故", "日程安排", "误车", "38人",
                        "认真核实", "措施不当", "已经误车",
                        "每题只有一个答案", "至少有2个答案",
                        "正确的打", "错误的打"]
        to_remove = []
        for para in doc.paragraphs:
            text = para.text.strip()
            # 跳过模板头部信息（课程名/院系等）
            if any(kw in text for kw in old_headers + old_keywords):
                to_remove.append(para)
                continue
            # 删除纯数字+字母的题号行、答案表头
            if re.match(r'^[一二三四五六七八九十\d]+\s*$', text) and len(text) <= 4:
                to_remove.append(para)
        for p in to_remove:
            try:
                p._element.getparent().remove(p._element)
            except:
                pass

        # 删除所有旧表格（我们从零建）
        for ti in range(len(doc.tables) - 1, -1, -1):
            try:
                doc.tables[ti]._tbl.getparent().remove(doc.tables[ti]._tbl)
            except:
                pass

        # 删除头部后的所有空段落（无论答案在哪里）
        try:
            body = doc.element.body
            children = list(body)
            # 找到最后一个非空段落
            last_nonempty = -1
            for i, child in enumerate(children):
                if child.tag != '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p':
                    continue
                texts = child.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                text = ''.join(t.text or '' for t in texts)
                if text.strip():
                    last_nonempty = i
            # 从最后往前删空段落
            if last_nonempty >= 0:
                for i in range(len(children) - 1, last_nonempty, -1):
                    child = list(body)[i]
                    if child.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p':
                        texts = child.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                        text = ''.join(t.text or '' for t in texts)
                        if not text.strip():
                            body.remove(child)
        except Exception as e:
            print(f"  清理空段落失败: {e}")

    def _make_answer_table(self, doc, title, qs, score, is_tf=False, q_start=0):
        """创建答案表格（全局编号）"""
        total_sec = len(qs) * score
        p = doc.add_paragraph()
        r = p.add_run(f"{title}（每题{score}分，共{total_sec}分）")
        r.bold = True

        from docx.oxml.ns import qn
        cols = 10
        n_qs = len(qs)
        n_rows = (n_qs + cols - 1) // cols
        table = doc.add_table(rows=n_rows * 2, cols=cols)
        tbl_pr = table._tbl.tblPr if table._tbl.tblPr is not None else OxmlElement('w:tblPr')
        tbl_borders = OxmlElement('w:tblBorders')
        for edge in ('top','left','bottom','right','insideH','insideV'):
            border = OxmlElement(f'w:{edge}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tbl_borders.append(border)
        tbl_pr.append(tbl_borders)

        for qi, q in enumerate(qs):
            ri = (qi // cols) * 2
            ci = qi % cols
            table.rows[ri].cells[ci].text = str(q_start + qi + 1)
            for pa in table.rows[ri].cells[ci].paragraphs:
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 答案行
            if is_tf:
                ans = q.get("answer", q.get("ans", ""))
                if isinstance(ans, bool):
                    table.rows[ri + 1].cells[ci].text = "√" if ans else "×"
                else:
                    s = str(ans).strip().lower()
                    table.rows[ri + 1].cells[ci].text = "√" if s in ("true", "正确") else "×"
            else:
                ans = q.get("ans", "")
                if isinstance(ans, int):
                    table.rows[ri + 1].cells[ci].text = ["A","B","C","D"][ans] if ans < 4 else str(ans)
                else:
                    table.rows[ri + 1].cells[ci].text = str(ans)
            for pa in table.rows[ri + 1].cells[ci].paragraphs:
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                pa.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _build_scoring_sections(self, doc, selected, sections):
        """根据实际选题构建评分标准答案部分"""
        KEY_TITLES = {
            "choice": "选择题", "multiple": "多选题", "tf": "判断题", "fill": "填空题",
            "short": "简答题", "calc": "计算分析题", "analysis": "分析题",
            "分析题": "分析题", "应用题": "应用题", "应用分析题": "应用分析题",
        }
        CN_NUMS = ["一", "二", "三", "四", "五", "六", "七", "八"]

        q_global = 0
        for si, (title, key, count, score) in enumerate(sections):
            qs = selected.get(key, [])
            if not qs:
                continue

            if not title.strip():
                cn = CN_NUMS[si] if si < len(CN_NUMS) else str(si + 1)
                type_name = KEY_TITLES.get(key, key)
                title = f"{cn}、{type_name}"
            total_sec = len(qs) * score

            # 选择题——用表格填答案
            if key in ("choice", "multiple"):
                self._make_answer_table(doc, title, qs, score, is_tf=False, q_start=q_global)
                q_global += len(qs)

            # 判断题——用表格填答案
            elif key == "tf":
                self._make_answer_table(doc, title, qs, score, is_tf=True, q_start=q_global)
                q_global += len(qs)

            # 其他题型——用段落列出答案
            else:
                p = doc.add_paragraph()
                r = p.add_run(f"{title}答案（每题{score}分，共{total_sec}分）")
                r.bold = True
                for i, q in enumerate(qs, 1):
                    q_global += 1
                    # 填空题：直接显示答案和分值，不拆得分点
                    if key == "fill":
                        ans = q.get("ans", q.get("answer", ""))
                        ans_text = str(ans) if str(ans) != "见解析" else ""
                        doc.add_paragraph(f"  {q_global}.  {ans_text} （{score}分）")
                        continue
                    
                    exp = q.get("exp", "")
                    cn_nums = ['①','②','③','④','⑤','⑥','⑦','⑧','⑨','⑩']
                    
                    if exp:
                        doc.add_paragraph(f"  {q_global}. （共{score}分）")
                        # 分割得分点
                        points = []
                        buf = ''
                        for ch in exp:
                            buf += ch
                            m = re.search(r'（\d+分）$', buf)
                            if m:
                                pre = buf[:m.start()]
                                if pre.strip():
                                    points.append(pre.strip())
                                points.append(m.group())
                                buf = ''
                        if buf.strip():
                            points.append(buf.strip())
                        merged = []
                        for p in points:
                            if re.match(r'^（\d+分）$', p.strip()) and merged:
                                merged[-1] += p
                            else:
                                merged.append(p)
                        points = merged
                        if not points or all(len(p) < 3 for p in points):
                            pts2 = re.split(r'(?<=\d分)(?=\S)', exp)
                            points = [p.strip() for p in pts2 if p.strip()]
                        if len(points) <= 1:
                            pts3 = re.findall(r'[^ ]+?\d+分', exp)
                            if pts3: points = pts3
                        # 按子问题（1）（2）分组合并
                        groups = []
                        cur_group = []
                        for pt in points:
                            if re.match(r'（\d+）', pt.strip()):
                                if cur_group:
                                    groups.append(' '.join(cur_group))
                                cur_group = [pt.strip()]
                            else:
                                cur_group.append(pt.strip())
                        if cur_group:
                            groups.append(' '.join(cur_group))
                        # 分组渲染
                        for gi, grp in enumerate(groups):
                            grp = grp.strip()
                            if not grp: continue
                            if not re.search(r'（\d+分）\s*$', grp):
                                grp = re.sub(r'(\d+)分\s*$', r'（\1分）', grp)
                            # 如果组已以（数字）开头，不再加①②
                            if re.match(r'（\d+）', grp):
                                doc.add_paragraph(f"    {grp}")
                            else:
                                marker = cn_nums[gi] if gi < len(cn_nums) else f'[{gi+1}]'
                                doc.add_paragraph(f"    {marker}{grp}")
                    else:
                        ans = q.get("ans", q.get("answer", ""))
                        ans_text = str(ans) if str(ans) != "见解析" else ""
                        explanation = q.get("explanation", "")
                        main_text = ans_text or explanation or ""
                        if main_text:
                            doc.add_paragraph(f"  {q_global}.  {main_text} （{score}分）")
                        else:
                            doc.add_paragraph(f"  {q_global}. （{score}分）")

    def _insert_figure(self, doc, q):
        """插入配图 — 多种来源"""
        import re
        from PIL import Image as PILImage

        def _try_insert(paths: list) -> bool:
            for p in paths:
                if os.path.exists(p):
                    try:
                        img = PILImage.open(p)
                        w, h = img.size
                        ratio = w / h
                        if ratio > 1.2:  # 横图
                            pw = Inches(6.0)
                        elif ratio < 0.8:  # 竖图
                            pw = Inches(4.0)
                        else:  # 方图
                            pw = Inches(4.0)
                        doc.add_picture(p, width=pw)
                        doc.add_paragraph("")
                        return True
                    except Exception as e:
                        print(f"  插图失败 {p}: {e}")
            return False

        # 1) imgs 字段（电工电子技术: "figures/图NEW-01.png"）
        imgs = q.get("imgs", q.get("image", ""))
        if imgs:
            if isinstance(imgs, list):
                for rel in imgs:
                    p = os.path.join(FIG_DIR_EE, os.path.basename(rel))
                    if _try_insert([p]): return
            else:
                # 单个字符串: "img/image-xxx.png" 或 "figures/图NEW-01.png"
                p1 = os.path.join(_BASE, "data", imgs)
                p2 = os.path.join(FIG_DIR_EE, os.path.basename(imgs))
                if _try_insert([p1, p2]): return

        # 2) figs 字段（电工电子技术: ["NEW-01"]）
        figs = q.get("figs", [])
        if figs:
            for f in figs:
                if _try_insert([os.path.join(FIG_DIR_EE, f"图{f}.png")]):
                    return

        # 3) 文本匹配 图N 或 NEW-N
        text = q.get("q", q.get("text", ""))
        m = re.search(r'(NEW-\d+)', text)
        if m:
            if _try_insert([os.path.join(FIG_DIR_EE, f"图{m.group(1)}.png")]):
                return
        m = re.search(r'图\s*(\d+)', text)
        if m:
            dirs = [FIG_DIR_EE, os.path.join(_BASE, "data", "img")]
            if _try_insert([os.path.join(d, f"图{m.group(1)}.png") for d in dirs]):
                return

"""验证答题区域和图片"""
from docx import Document

d = Document('output/单片机技术_期末试卷_A.docx')
print(f'A卷: {len(d.paragraphs)}段')

for i, p in enumerate(d.paragraphs):
    t = p.text.strip()
    has_img = 'w:drawing' in p._element.xml
    m = ' [📷]' if has_img else ''
    if t.startswith(('26.','27.','28.','29.')) or has_img:
        print(f'  P[{i:3d}]: {t[:55]}{m}')

# 统计空行数
in_q = None
blanks = 0
for i, p in enumerate(d.paragraphs):
    t = p.text.strip()
    if t.startswith('26.'):
        in_q = 26; blanks = 0
    elif t.startswith('27.'):
        if in_q == 26: print(f'  Q26→Q27: {blanks}空行')
        in_q = 27; blanks = 0
    elif t.startswith('28.'):
        if in_q == 27: print(f'  Q27→Q28: {blanks}空行')
        in_q = 28; blanks = 0
    elif t.startswith('29.'):
        if in_q == 28: print(f'  Q28→Q29: {blanks}空行')
        in_q = 29; blanks = 0
    elif not t:
        blanks += 1
if in_q == 29: print(f'  Q29→END: {blanks}空行')

# B卷
d = Document('output/单片机技术_期末试卷_B.docx')
print(f'\nB卷: {len(d.paragraphs)}段')
in_q = None
blanks = 0
for i, p in enumerate(d.paragraphs):
    t = p.text.strip()
    if t.startswith('26.'):
        in_q = 26; blanks = 0
    elif t.startswith('27.'):
        if in_q == 26: print(f'  Q26→Q27: {blanks}空行')
        in_q = 27; blanks = 0
    elif t.startswith('28.'):
        if in_q == 27: print(f'  Q27→Q28: {blanks}空行')
        in_q = 28; blanks = 0
    elif t.startswith('29.'):
        if in_q == 28: print(f'  Q28→Q29: {blanks}空行')
        in_q = 29; blanks = 0
    elif not t:
        blanks += 1
if in_q == 29: print(f'  Q29→END: {blanks}空行')

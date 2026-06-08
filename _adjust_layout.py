"""调整单片机试卷答题区域和图片尺寸"""
from docx import Document
from docx.shared import Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree
from io import BytesIO

# ================================================================
# 1. 处理A卷试卷
# ================================================================
print("=== 单片机技术 A卷试卷 ===")
doc = Document('output/单片机技术_期末试卷_A.docx')
paras = doc.paragraphs
body = doc.element.body

# 分析简答题和综合题位置
print("简答/综合题区域:")
for i, p in enumerate(paras):
    t = p.text.strip()
    if '四、简答题' in t or '五、综合应用题' in t or '简答题' in t or '综合应用' in t:
        print(f"  P[{i}]: {t[:50]}")
    if t.startswith(('26.', '27.', '28.', '29.')):
        print(f"  P[{i}]: {t[:60]}")

# 找到综合题附近是否有图片
print("\n图片位置:")
for i, p in enumerate(paras):
    if 'w:drawing' in p._element.xml:
        print(f"  P[{i}]: 含图片")

# 增大简答题 Q26, Q27 的答题区域
# 在 Q26(10分) 后增加更多空行 - 从6行增加到18行
# 在 Q27(10分) 后增加更多空行
# 在综合题 Q28, Q29 后增加更多空行

# 找到四、简答题位置
short_title_idx = None
for i, p in enumerate(paras):
    if '四、简答题' in p.text:
        short_title_idx = i
        break

if short_title_idx:
    print(f"\n简答题标题在 P[{short_title_idx}]")
    # 找到Q26和Q27的段落
    q26_idx = None
    q27_idx = None
    for i in range(short_title_idx, len(paras)):
        t = paras[i].text.strip()
        if t.startswith('26.') and q26_idx is None:
            q26_idx = i
        elif t.startswith('27.') and q26_idx is not None and q27_idx is None:
            q27_idx = i
            break
    
    print(f"Q26在 P[{q26_idx}], Q27在 P[{q27_idx}]")
    
    # 在Q26后加12个空行（已有6个，再加12）
    if q26_idx:
        body_elem = doc.element.body
        q26_elem = paras[q26_idx]._element
        q26_pos = list(body_elem).index(q26_elem)
        # 找到Q27位置
        q27_elem = paras[q27_idx]._element
        q27_pos = list(body_elem).index(q27_elem)
        
        # 在Q26和Q27之间插入空行
        for n in range(12):
            p = etree.SubElement(body_elem, qn('w:p'))
            body_elem.insert(q27_pos + n, p)
        
        # 重新获取段落索引
        doc.save('output/单片机技术_期末试卷_A_temp.docx')
        doc = Document('output/单片机技术_期末试卷_A_temp.docx')
        paras = doc.paragraphs
        body = doc.element.body
        print("  在Q26和Q27之间插入12个空行")

# 重新查找位置
for i, p in enumerate(paras):
    t = p.text.strip()
    if t.startswith('27.'):
        q27_idx = i
    if t.startswith('28.'):
        q28_idx = i
    if t.startswith('29.'):
        q29_idx = i

# 在Q27后加12个空行
if q27_idx:
    body_elem = doc.element.body
    q28_idx = None
    for i, p in enumerate(paras):
        if p.text.strip().startswith('28.'):
            q28_idx = i
            break
    
    if q28_idx:
        q27_elem = paras[q27_idx]._element
        for n in range(12):
            p = etree.SubElement(body_elem, qn('w:p'))
            body_elem.insert(q28_idx - 1, p)  # 在Q28前插入
        
        doc.save('output/单片机技术_期末试卷_A_temp.docx')
        doc = Document('output/单片机技术_期末试卷_A_temp.docx')
        paras = doc.paragraphs
        body = doc.element.body
        print("  在Q27和Q28之间插入12个空行")

# 重新查找Q28和Q29位置
q28_idx = None
q29_idx = None
analysis_end_idx = None
for i, p in enumerate(paras):
    t = p.text.strip()
    if t.startswith('28.'):
        q28_idx = i
    elif t.startswith('29.'):
        q29_idx = i
    elif '五、综合应用题' in t:
        analysis_title_idx = i

if q28_idx and q29_idx:
    print(f"Q28在 P[{q28_idx}], Q29在 P[{q29_idx}]")
    body_elem = doc.element.body
    
    # Q28后加12个空行（在Q29前）
    q28_elem = paras[q28_idx]._element
    q29_elem = paras[q29_idx]._element
    q29_pos = list(body_elem).index(q29_elem)
    
    for n in range(12):
        p = etree.SubElement(body_elem, qn('w:p'))
        body_elem.insert(q29_pos + n, p)
    
    doc.save('output/单片机技术_期末试卷_A_temp.docx')
    doc = Document('output/单片机技术_期末试卷_A_temp.docx')
    paras = doc.paragraphs
    body = doc.element.body
    print("  在Q28和Q29之间插入12个空行")

# Q29后加12个空行（在文末）
q29_idx = None
for i, p in enumerate(paras):
    if p.text.strip().startswith('29.'):
        q29_idx = i
        break

if q29_idx:
    body_elem = doc.element.body
    q29_elem = paras[q29_idx]._element
    q29_pos = list(body_elem).index(q29_elem)
    
    for n in range(12):
        p = etree.SubElement(body_elem, qn('w:p'))
        body_elem.insert(q29_pos + 1 + n, p)
    
    print("  在Q29后插入12个空行")

# ================================================================
# 2. 放大图片
# ================================================================
print("\n放大图片...")
image_parts = {}
for rel in doc.part.rels.values():
    if "image" in rel.reltype:
        image_parts[rel.rId] = rel.target_part.blob

# 查找并放大所有图片
for i, p in enumerate(doc.paragraphs):
    xml = p._element.xml
    if 'w:drawing' in xml:
        # 获取当前图片尺寸
        root = etree.fromstring(xml.encode())
        ns = {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
              'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}
        
        for ext in root.iter('{http://schemas.openxmlformats.org/drawingml/2006/main}ext'):
            cx = ext.get('cx')
            cy = ext.get('cy')
            if cx and cy:
                old_w = int(cx) / 914400
                old_h = int(cy) / 914400
                print(f"  P[{i}] 图片: {old_w:.1f}in x {old_h:.1f}in")
                
                # 放大到新的尺寸
                new_w = old_w * 1.6
                new_h = old_h * 1.6
                ext.set('cx', str(int(new_w * 914400)))
                ext.set('cy', str(int(new_h * 914400)))
                print(f"    → 放大到 {new_w:.1f}in x {new_h:.1f}in")
        
        # 同时更新wp:extent
        for wp_ext in root.iter('{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent'):
            cx = wp_ext.get('cx')
            cy = wp_ext.get('cy')
            if cx and cy:
                old_w = int(cx) / 914400
                old_h = int(cy) / 914400
                new_w = old_w * 1.6
                new_h = old_h * 1.6
                wp_ext.set('cx', str(int(new_w * 914400)))
                wp_ext.set('cy', str(int(new_h * 914400)))
        
        # 写回修改后的XML
        p._element.clear()
        for child in root:
            p._element.append(child)

import os as _os
if _os.path.exists('output/单片机技术_期末试卷_A_temp.docx'):
    _os.remove('output/单片机技术_期末试卷_A_temp.docx')

doc.save('output/单片机技术_期末试卷_A.docx')
print("\n✅ A卷已保存")

# ================================================================
# 3. 处理B卷（B卷是A卷的副本，同样处理）
# ================================================================
print("\n=== 单片机技术 B卷试卷 ===")
# B卷没有图片，只需增加答题区域
doc = Document('output/单片机技术_期末试卷_B.docx')
paras = doc.paragraphs
body = doc.element.body

# 同样在简答和综合题后加空行
short_title_idx = None
for i, p in enumerate(paras):
    if '四、简答题' in p.text:
        short_title_idx = i
        break

if short_title_idx:
    q26_idx = None
    q27_idx = None
    for i in range(short_title_idx, len(paras)):
        t = paras[i].text.strip()
        if t.startswith('26.') and q26_idx is None:
            q26_idx = i
        elif t.startswith('27.') and q26_idx is not None and q27_idx is None:
            q27_idx = i
            break
    
    # Q26后加12空行
    q28_idx = None
    for i, p in enumerate(paras):
        if p.text.strip().startswith('28.'):
            q28_idx = i
            break
    
    if q26_idx and q27_idx:
        q26_elem = paras[q26_idx]._element
        for n in range(12):
            p = etree.SubElement(body, qn('w:p'))
            body.insert(q27_idx - 1 + n, p)
        
        doc.save('output/单片机技术_期末试卷_B_temp.docx')
        doc = Document('output/单片机技术_期末试卷_B_temp.docx')
        paras = doc.paragraphs
        body = doc.element.body
    
    # Q27后加12空行
    q27_idx = None
    q28_idx = None
    for i, p in enumerate(paras):
        t = p.text.strip()
        if t.startswith('27.'): q27_idx = i
        if t.startswith('28.'): q28_idx = i
    
    if q27_idx and q28_idx:
        q27_elem = paras[q27_idx]._element
        for n in range(12):
            p = etree.SubElement(body, qn('w:p'))
            body.insert(q28_idx - 1 + n, p)
        
        doc.save('output/单片机技术_期末试卷_B_temp.docx')
        doc = Document('output/单片机技术_期末试卷_B_temp.docx')
        paras = doc.paragraphs
        body = doc.element.body
    
    # Q28后加12空行
    q28_idx = None
    q29_idx = None
    for i, p in enumerate(paras):
        t = p.text.strip()
        if t.startswith('28.'): q28_idx = i
        if t.startswith('29.'): q29_idx = i
    
    if q28_idx and q29_idx:
        q28_elem = paras[q28_idx]._element
        for n in range(12):
            p = etree.SubElement(body, qn('w:p'))
            body.insert(q29_idx - 1 + n, p)
        
        doc.save('output/单片机技术_期末试卷_B_temp.docx')
        doc = Document('output/单片机技术_期末试卷_B_temp.docx')
        paras = doc.paragraphs
        body = doc.element.body
    
    # Q29后加12空行
    q29_idx = None
    for i, p in enumerate(paras):
        if p.text.strip().startswith('29.'):
            q29_idx = i
            break
    
    if q29_idx:
        q29_elem = paras[q29_idx]._element
        for n in range(12):
            p = etree.SubElement(body, qn('w:p'))
            body.insert(q29_idx + 1 + n, p)
        
        print("B卷: 答题区域已扩充")

if _os.path.exists('output/单片机技术_期末试卷_B_temp.docx'):
    _os.remove('output/单片机技术_期末试卷_B_temp.docx')

doc.save('output/单片机技术_期末试卷_B.docx')
print("✅ B卷已保存")

print("\n全部完成")
print("单片机A卷: 答题区域扩充+图片放大1.6x")
print("单片机B卷: 答题区域扩充")
